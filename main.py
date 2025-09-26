import os
import sys 
from fpdf import FPDF
from docx import Document
from PyPDF2 import PdfReader


def converter_txt_para_pdf(caminho_txt, caminho_pdf):

    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        

        with open(caminho_txt, 'r', encoding='utf-8') as arquivo_txt:
            for linha in arquivo_txt:

                pdf.multi_cell(0, 10, linha) 
        
        pdf.output(caminho_pdf)
        print(f"\n[SUCESSO] Arquivo PDF salvo em: {caminho_pdf}")
    except UnicodeDecodeError:
        print("\n[ERRO] Problema de codificação ao ler o arquivo TXT. Certifique-se de que está em UTF-8.")
    except Exception as e:
        print(f"\n[ERRO] Ocorreu um erro durante a conversão: {e}")

def converter_txt_para_docx(caminho_txt, caminho_docx):

    try:
        document = Document()
        with open(caminho_txt, 'r', encoding='utf-8') as arquivo_txt:
            for linha in arquivo_txt:
                document.add_paragraph(linha.strip()) 
        
        document.save(caminho_docx)
        print(f"\n[SUCESSO] Arquivo docx salvo em: {caminho_docx}")
    except Exception as e:
        print(f"\n[ERRO] Ocorreu um erro durante a conversão: {e}")

def converter_pdf_para_docx(caminho_pdf, caminho_docx):

    try:
        document = Document()
        with open(caminho_pdf, 'rb') as arquivo_pdf:
            leitor_pdf = PdfReader(arquivo_pdf)
            for pagina in leitor_pdf.pages:
                texto = pagina.extract_text()
                if texto:
                    document.add_paragraph(texto)
        
        document.save(caminho_docx)
        print(f"\n[SUCESSO] Arquivo docx salvo em: {caminho_docx}")
    except Exception as e:
        print(f"\n[ERRO] Ocorreu um erro durante a conversão: {e}")

def converter_docx_para_pdf(caminho_docx, caminho_pdf):

    try:
        document = Document(caminho_docx)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        
        for paragrafo in document.paragraphs:
            pdf.multi_cell(0, 10, paragrafo.text)
        
        pdf.output(caminho_pdf)
        print(f"\n[SUCESSO] Arquivo PDF salvo em: {caminho_pdf}")
    except Exception as e:
        print(f"\n[ERRO] Ocorreu um erro durante a conversão: {e}")



def obter_caminho_entrada(extensao, tipo):

    while True:
        caminho = input(f"Digite o caminho do arquivo {tipo} ({extensao}): ")
        if os.path.isfile(caminho):
            return caminho
        print(f"[ERRO] Arquivo {tipo} não encontrado. Tente novamente.")

def obter_caminho_saida(extensao, tipo):

    while True:
        caminho = input(f"Digite o caminho onde o arquivo {tipo} será salvo (ex: nome.{extensao}): ")
        if caminho.lower().endswith(f'.{extensao}'):
            return caminho
        print(f"[ERRO] O caminho do arquivo deve terminar com .{extensao}. Tente novamente.")


def main():
    menu = (
        "1 - Converter de txt para PDF\n"
        "2 - Converter de txt para docx\n"
        "3 - Converter de PDF para docx\n"
        "4 - Converter de docx para PDF\n"
        "Escolha o número do conversor desejado: "
    )
    escolha_do_tipo_de_conversor = input(menu).strip()

    if escolha_do_tipo_de_conversor == "1":
        caminho_in = obter_caminho_entrada('txt', 'txt')
        caminho_out = obter_caminho_saida('pdf', 'PDF')
        converter_txt_para_pdf(caminho_in, caminho_out)
        
    elif escolha_do_tipo_de_conversor == "2":
        caminho_in = obter_caminho_entrada('txt', 'txt')
        caminho_out = obter_caminho_saida('docx', 'docx')
        converter_txt_para_docx(caminho_in, caminho_out)
        
    elif escolha_do_tipo_de_conversor == "3":
        caminho_in = obter_caminho_entrada('pdf', 'PDF')
        caminho_out = obter_caminho_saida('docx', 'docx')

        converter_pdf_para_docx(caminho_in, caminho_out)
        
    elif escolha_do_tipo_de_conversor == "4":
        caminho_in = obter_caminho_entrada('docx', 'docx')
        caminho_out = obter_caminho_saida('pdf', 'PDF')

        converter_docx_para_pdf(caminho_in, caminho_out)
        
    else:
        print("\n[ERRO] Opção inválida. O programa será encerrado.")

        sys.exit(1)
        
    print("\nPressione Enter para sair...")
    input() 

if __name__ == "__main__":
    main()